import streamlit as st
import os
from typing import List, Dict, Any
import tempfile
import uuid
from datetime import datetime

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.schema import Document
from langchain_groq import ChatGroq
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate

from crewai import Agent, Task, Crew, Process

import PyPDF2
import docx
import json
from io import BytesIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor

st.set_page_config(
    page_title="AI Resume Summarizer Agent",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.5rem;
        color: #2c3e50;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .summary-card {
        background-color: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 4px solid #1f77b4;
        margin: 1rem 0;
    }
    .agent-output {
        background-color: #ffffff;
        padding: 1.5rem;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .chat-container {
        background-color: #ffffff;
        padding: 1rem;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
        margin: 1rem 0;
    }
    .file-upload-area {
        border: 2px dashed #cccccc;
        border-radius: 10px;
        padding: 2rem;
        text-align: center;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state variables
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = ""
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {}
if 'conversation_chain' not in st.session_state:
    st.session_state.conversation_chain = None

class ResumeProcessor:
    def __init__(self):
        self.groq_api_key = os.getenv("GROQ_API_KEY")
        
        self.llm = ChatGroq(
            temperature=0.1,
            groq_api_key=self.groq_api_key,
            model_name="llama3-8b-8192"
        )
        
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
    
    def extract_text_from_file(self, uploaded_file) -> str:
        try:
            if uploaded_file.type == "application/pdf":
                return self._extract_from_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_from_docx(uploaded_file)
            elif uploaded_file.type == "text/plain":
                return str(uploaded_file.read(), "utf-8")
            else:
                st.error("Unsupported file format. Please upload PDF, DOCX, or TXT files.")
                return ""
        except Exception as e:
            st.error(f"Error extracting text: {str(e)}")
            return ""
    
    def _extract_from_pdf(self, pdf_file) -> str:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    
    def _extract_from_docx(self, docx_file) -> str:
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def create_vector_store(self, text: str):
        documents = [Document(page_content=text)]
        chunks = self.text_splitter.split_documents(documents)
        vector_store = FAISS.from_documents(chunks, self.embeddings)
        return vector_store
    
    def setup_conversation_chain(self, vector_store):
        memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="answer"
        )
        
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=self.llm,
            retriever=vector_store.as_retriever(search_kwargs={"k": 3}),
            memory=memory,
            return_source_documents=True
        )
        
        return conversation_chain
    
    def create_crew_agents(self, resume_text: str):
        crew_llm = ChatGroq(
            temperature=0.1,
            groq_api_key=self.groq_api_key,
            model_name="groq/llama3-8b-8192"  # Fixed model name
        )
        
        resume_analyzer = Agent(
            role="Senior Resume Analyst",
            goal="Analyze resumes comprehensively to extract key information and insights",
            backstory="You are an expert HR professional with 15+ years of experience in talent acquisition and resume analysis. You have a keen eye for identifying candidate strengths, potential red flags, and cultural fit indicators.",
            verbose=True,
            allow_delegation=False,
            llm=crew_llm
        )
        
        skills_evaluator = Agent(
            role="Technical Skills Evaluator",
            goal="Evaluate and categorize technical and soft skills from resumes",
            backstory="You are a technical recruiter specializing in evaluating candidate skills across various domains including technology, management, and soft skills. You understand market demand and skill relevance.",
            verbose=True,
            allow_delegation=False,
            llm=crew_llm
        )
        
        market_researcher = Agent(
            role="Market Research Specialist",
            goal="Research current market trends and salary benchmarks for candidates",
            backstory="You are a market research expert who stays updated on industry trends, salary benchmarks, and hiring patterns. You provide valuable insights on candidate market positioning.",
            verbose=True,
            allow_delegation=False,
            llm=crew_llm
        )
        
        return resume_analyzer, skills_evaluator, market_researcher
    
    def create_crew_tasks(self, resume_text: str, resume_analyzer, skills_evaluator, market_researcher):
        analysis_task = Task(
            description=f"""
            Analyze the following resume comprehensively and provide a detailed structured report:
            
            {resume_text}
            
            Create a comprehensive analysis with the following sections:
            
            ## CANDIDATE PROFILE SUMMARY
            - Professional summary and career level
            - Years of experience and expertise areas
            - Current role and responsibilities
            
            ## WORK EXPERIENCE ANALYSIS
            - Detailed breakdown of each role
            - Key responsibilities and achievements
            - Career progression and growth trajectory
            - Notable projects and accomplishments
            
            ## EDUCATION & CERTIFICATIONS
            - Educational background analysis
            - Relevant certifications and training
            - Academic achievements
            
            ## KEY STRENGTHS & ACHIEVEMENTS
            - Top 5 professional strengths
            - Quantifiable achievements and results
            - Leadership and teamwork examples
            
            ## AREAS FOR IMPROVEMENT
            - Skills gaps identified
            - Experience limitations
            - Recommendations for professional development
            
            ## OVERALL ASSESSMENT
            - Suitability for target roles
            - Competitive positioning
            - Final recommendation and next steps
            
            Format the response with clear headers and bullet points for easy reading.
            """,
            agent=resume_analyzer,
            expected_output="A comprehensive 500+ word resume analysis report with structured sections and detailed insights"
        )
        
        skills_task = Task(
            description=f"""
            Conduct a thorough skills evaluation of this resume:
            
            {resume_text}
            
            Provide a detailed skills assessment with these sections:
            
            ## TECHNICAL SKILLS BREAKDOWN
            - Programming languages and frameworks
            - Tools and technologies
            - Software and platforms
            - Skill proficiency levels (Beginner/Intermediate/Advanced)
            
            ## SOFT SKILLS IDENTIFICATION
            - Communication and collaboration skills
            - Leadership and management abilities
            - Problem-solving and analytical skills
            - Adaptability and learning agility
            
            ## INDUSTRY-SPECIFIC SKILLS
            - Domain expertise and specializations
            - Industry knowledge and experience
            - Compliance and regulatory awareness
            
            ## SKILL RELEVANCE ANALYSIS
            - Current market demand for these skills
            - Emerging technologies alignment
            - Skills transferability across industries
            
            ## SKILL GAPS & DEVELOPMENT RECOMMENDATIONS
            - Missing skills for career advancement
            - Trending skills to acquire
            - Learning path suggestions
            - Certification recommendations
            
            Provide specific examples from the resume to support your analysis.
            """,
            agent=skills_evaluator,
            expected_output="A detailed 400+ word skills evaluation report with categorized skills and specific recommendations"
        )
        
        market_task = Task(
            description=f"""
            Conduct market research and positioning analysis for this candidate:
            
            {resume_text}
            
            Provide a comprehensive market analysis with these sections:
            
            ## MARKET DEMAND ANALYSIS
            - Current job market trends for this profile
            - High-demand roles and positions
            - Industry growth sectors
            - Geographic market opportunities
            
            ## SALARY BENCHMARKING
            - Estimated salary range based on experience
            - Industry compensation standards
            - Benefits and perks expectations
            - Negotiation leverage points
            
            ## COMPETITIVE LANDSCAPE
            - Candidate's market positioning
            - Competitive advantages
            - Unique value propositions
            - Differentiating factors
            
            ## INDUSTRY TRENDS & OPPORTUNITIES
            - Emerging technologies and their impact
            - Future skill requirements
            - Career path opportunities
            - Industry challenges and solutions
            
            ## STRATEGIC RECOMMENDATIONS
            - Target companies and roles
            - Networking strategies
            - Personal branding suggestions
            - Career advancement roadmap
            
            Base your analysis on general industry knowledge and current market conditions.
            """,
            agent=market_researcher,
            expected_output="A comprehensive 400+ word market research report with salary insights and strategic recommendations"
        )
        
        return analysis_task, skills_task, market_task
    
    def run_crew_analysis(self, resume_text: str):
        try:
            resume_analyzer, skills_evaluator, market_researcher = self.create_crew_agents(resume_text)
            
            analysis_task, skills_task, market_task = self.create_crew_tasks(
                resume_text, resume_analyzer, skills_evaluator, market_researcher
            )
            
            crew = Crew(
                agents=[resume_analyzer, skills_evaluator, market_researcher],
                tasks=[analysis_task, skills_task, market_task],
                process=Process.sequential,
                verbose=True
            )
            
            # Execute the crew and get results
            result = crew.kickoff()
            
            # Extract individual task results
            analysis_results = {
                'resume_analysis': str(analysis_task.output) if hasattr(analysis_task, 'output') else "Analysis not completed",
                'skills_evaluation': str(skills_task.output) if hasattr(skills_task, 'output') else "Skills evaluation not completed", 
                'market_research': str(market_task.output) if hasattr(market_task, 'output') else "Market research not completed",
                'combined_result': str(result)
            }
            
            return analysis_results
            
        except Exception as e:
            st.error(f"Error in crew analysis: {str(e)}")
            return None

    def generate_pdf_report(self, analysis_results, chat_history):
        """Generate PDF report from analysis results"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=HexColor('#1f77b4'),
            alignment=1
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=HexColor('#2c3e50'),
            spaceBefore=20
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=14,
            spaceAfter=10,
            textColor=HexColor('#34495e'),
            spaceBefore=15
        )
        
        story = []
        
        # Title
        story.append(Paragraph("🤖 AI Resume Analysis Report", title_style))
        story.append(Spacer(1, 12))
        
        # Generation info
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Resume Analysis Section
        if 'resume_analysis' in analysis_results:
            story.append(Paragraph("📊 Resume Analysis", heading_style))
            analysis_text = analysis_results['resume_analysis'].replace('\n', '<br/>')
            story.append(Paragraph(analysis_text, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Skills Evaluation Section  
        if 'skills_evaluation' in analysis_results:
            story.append(Paragraph("🎯 Skills Evaluation", heading_style))
            skills_text = analysis_results['skills_evaluation'].replace('\n', '<br/>')
            story.append(Paragraph(skills_text, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Market Research Section
        if 'market_research' in analysis_results:
            story.append(Paragraph("📈 Market Research", heading_style))
            market_text = analysis_results['market_research'].replace('\n', '<br/>')
            story.append(Paragraph(market_text, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Chat History Section
        if chat_history:
            story.append(PageBreak())
            story.append(Paragraph("💬 Chat History", heading_style))
            for i, chat in enumerate(chat_history):
                story.append(Paragraph(f"<b>Question {i+1}:</b> {chat['question']}", styles['Normal']))
                story.append(Spacer(1, 6))
                story.append(Paragraph(f"<b>Answer:</b> {chat['answer']}", styles['Normal']))
                story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        return buffer

def main():
    st.markdown('<h1 class="main-header">🤖 AI Resume Summarizer Agent</h1>', unsafe_allow_html=True)
    
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        groq_api_key = st.text_input("Groq API Key", type="password", value=os.getenv("GROQ_API_KEY", ""))
        
        if groq_api_key:
            os.environ["GROQ_API_KEY"] = groq_api_key
        
        st.divider()
        
        st.header("📋 Instructions")
        st.markdown("""
        1. **Add API Key** above
        2. **Upload Resume** (PDF, DOCX, TXT)
        3. **Get AI Analysis** using CrewAI
        4. **Chat with Resume** using RAG
        """)
        
        st.divider()
        
        if st.button("🗑️ Clear Chat History"):
            st.session_state.chat_history = []
            st.rerun()
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown('<h2 class="section-header">📁 Upload Resume</h2>', unsafe_allow_html=True)
        
        uploaded_file = st.file_uploader(
            "Choose a resume file",
            type=['pdf', 'docx', 'txt'],
            help="Upload PDF, DOCX, or TXT files"
        )
        
        if uploaded_file is not None:
            if not groq_api_key:
                st.error("Please provide Groq API Key in the sidebar")
                return
            
            processor = ResumeProcessor()
            
            with st.spinner("Extracting text from resume..."):
                resume_text = processor.extract_text_from_file(uploaded_file)
                st.session_state.resume_text = resume_text
            
            if resume_text:
                st.success("✅ Resume text extracted successfully!")
                with st.expander("📄 Resume Text Preview"):
                    st.text_area("Resume Content", resume_text[:1000] + "..." if len(resume_text) > 1000 else resume_text, height=200)
                
                with st.spinner("Creating vector store for RAG..."):
                    vector_store = processor.create_vector_store(resume_text)
                    st.session_state.vector_store = vector_store
                    
                    conversation_chain = processor.setup_conversation_chain(vector_store)
                    st.session_state.conversation_chain = conversation_chain
                
                if st.button("🚀 Generate AI Analysis", type="primary"):
                    with st.spinner("Running CrewAI analysis... This may take a few minutes..."):
                        results = processor.run_crew_analysis(resume_text)
                        if results:
                            st.session_state.analysis_results = results
                            st.success("✅ Analysis complete!")
    
    with col2:
        st.markdown('<h2 class="section-header">💬 Chat with Resume</h2>', unsafe_allow_html=True)
        
        if st.session_state.conversation_chain:
            user_question = st.text_input("Ask something about the resume:", key="user_input")
            
            if user_question:
                with st.spinner("Thinking..."):
                    try:
                        response = st.session_state.conversation_chain({
                            "question": user_question
                        })
                        
                        st.session_state.chat_history.append({
                            "question": user_question,
                            "answer": response["answer"],
                            "timestamp": datetime.now().strftime("%H:%M:%S")
                        })
                        
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
            
            if st.session_state.chat_history:
                st.markdown("### 💬 Chat History")
                for i, chat in enumerate(reversed(st.session_state.chat_history)):
                    with st.container():
                        st.markdown(f"**🙋 Question ({chat['timestamp']}):** {chat['question']}")
                        st.markdown(f"**🤖 Answer:** {chat['answer']}")
                        st.divider()
        else:
            st.info("Upload a resume first to start chatting!")
    
    # Display all analysis results
    if st.session_state.analysis_results:
        st.markdown('<h2 class="section-header">📊 AI Analysis Results</h2>', unsafe_allow_html=True)
        
        # Create tabs for different analyses
        tab1, tab2, tab3 = st.tabs(["📊 Resume Analysis", "🎯 Skills Evaluation", "📈 Market Research"])
        
        with tab1:
            if 'resume_analysis' in st.session_state.analysis_results:
                st.markdown('<div class="agent-output">', unsafe_allow_html=True)
                st.markdown("### 📊 Resume Analysis by Senior Resume Analyst")
                st.markdown(st.session_state.analysis_results['resume_analysis'])
                st.markdown('</div>', unsafe_allow_html=True)
        
        with tab2:
            if 'skills_evaluation' in st.session_state.analysis_results:
                st.markdown('<div class="agent-output">', unsafe_allow_html=True)
                st.markdown("### 🎯 Skills Evaluation by Technical Skills Evaluator")
                st.markdown(st.session_state.analysis_results['skills_evaluation'])
                st.markdown('</div>', unsafe_allow_html=True)
        
        with tab3:
            if 'market_research' in st.session_state.analysis_results:
                st.markdown('<div class="agent-output">', unsafe_allow_html=True)
                st.markdown("### 📈 Market Research by Market Research Specialist")
                st.markdown(st.session_state.analysis_results['market_research'])
                st.markdown('</div>', unsafe_allow_html=True)
        
        # Download button
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 1, 1])
        with col2:
            if st.button("📥 Download Complete Analysis Report", type="primary"):
                processor = ResumeProcessor()
                
                pdf_buffer = processor.generate_pdf_report(
                    st.session_state.analysis_results,
                    st.session_state.chat_history
                )
                
                st.download_button(
                    label="📄 Download PDF Report",
                    data=pdf_buffer,
                    file_name=f"resume_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf"
                )

if __name__ == "__main__":
    main()